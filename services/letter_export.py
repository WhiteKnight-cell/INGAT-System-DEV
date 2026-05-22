from io import BytesIO

from docx import Document
from fpdf import FPDF


def _safe_pdf_text(text):
    """FPDF core fonts are Latin-1; replace unsupported characters."""
    return text.encode('latin-1', errors='replace').decode('latin-1')


def build_letter_pdf(letter_text, complaint_id):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font('Helvetica', 'B', 14)
    pdf.cell(0, 10, _safe_pdf_text('INGAT — Formal Complaint Letter'), ln=True)
    pdf.set_font('Helvetica', '', 10)
    pdf.cell(0, 8, _safe_pdf_text(f'Reference: #ING-{complaint_id:04d}'), ln=True)
    pdf.ln(4)
    pdf.set_font('Helvetica', '', 11)

    for paragraph in letter_text.split('\n'):
        line = paragraph.strip()
        if line:
            pdf.multi_cell(0, 6, _safe_pdf_text(line))
            pdf.ln(2)
        else:
            pdf.ln(4)

    raw = pdf.output()
    buffer = BytesIO(raw if isinstance(raw, (bytes, bytearray)) else bytes(raw))
    buffer.seek(0)
    return buffer


def build_letter_docx(letter_text, complaint_id):
    doc = Document()
    doc.add_heading('INGAT — Formal Complaint Letter', level=1)
    doc.add_paragraph(f'Reference: #ING-{complaint_id:04d}')

    for paragraph in letter_text.split('\n'):
        line = paragraph.strip()
        if line:
            doc.add_paragraph(line)
        else:
            doc.add_paragraph('')

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
