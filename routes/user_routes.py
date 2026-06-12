from flask import Blueprint, render_template, redirect, url_for, flash, request, abort, send_file
from utils import hash_password, verify_password, is_bcrypt_hash
from flask_login import login_user, logout_user, current_user
from datetime import datetime, timedelta
import re

from routes.auth import member_required
from utils import (
    generate_reset_token,
    send_email,
    verify_reset_token,
    generate_otp_code,
    send_verification_email,
)

user_bp = Blueprint('user', __name__, url_prefix='/user')

# seconds to wait before allowing resend
RESEND_COOLDOWN_SECONDS = 60

ALLOWED_VIOLATION_TYPES = [
    'Illegal Dumping',
    'Air Pollution',
    'Water Pollution',
    'Illegal Logging',
    'Others',
]


def _password_error(password):
    if len(password) < 8:
        return 'Password must be at least 8 characters.'
    if not re.search(r'[A-Z]', password):
        return 'Password must contain at least one uppercase letter.'
    if not re.search(r'[a-z]', password):
        return 'Password must contain at least one lowercase letter.'
    if not re.search(r'[0-9]', password):
        return 'Password must contain at least one number.'
    if not re.search(r'[!@#$%^&*(),.?":{}|<>]', password):
        return 'Password must contain at least one special character.'
    return None


def _create_email_verification(user):
    from extensions import db
    from models import EmailVerification

    otp_code = generate_otp_code()
    expires_at = datetime.utcnow() + timedelta(minutes=15)
    verification = EmailVerification(
        user_id=user.id,
        otp_code=otp_code,
        expires_at=expires_at,
    )
    db.session.add(verification)
    db.session.commit()
    return verification


def _send_account_verification_email(user, otp_code):
    sent = send_verification_email(user.email, user.full_name, otp_code)
    if not sent:
        print(f'Failed to send verification email for user {user.email}')
        print(f'Local OTP for {user.email}: {otp_code}')
    return sent


def _complaint_form_from_request():
    """Preserve submitted values when re-rendering the form after validation errors."""
    return {
        'violation_type': request.form.get('violation_type', '').strip(),
        'street_address': request.form.get('street_address', '').strip(),
        'barangay': request.form.get('barangay', '').strip(),
        'municipality': request.form.get('municipality', '').strip(),
        'date_incident': request.form.get('date_incident', '').strip(),
        'description': request.form.get('description', ''),
    }


def _render_submit_form(form=None):
    return render_template('user/submit_complaint.html', form=form or {})


def _register_form_from_request():
    """Preserve submitted values when re-rendering registration after validation errors."""
    return {
        'full_name': request.form.get('full_name', '').strip(),
        'email': request.form.get('email', '').strip(),
        'contact_number': request.form.get('contact_number', '').strip(),
        'barangay': request.form.get('barangay', '').strip(),
        'municipality': request.form.get('municipality', '').strip(),
        'password': request.form.get('password', ''),
        'confirm_password': request.form.get('confirm_password', ''),
    }


def _render_register_form(form=None):
    return render_template('user/register.html', form=form or {}, f=form or {})


def _generate_letter_for_complaint(complaint, complainant):
    """Generate and persist Gemini letter. Returns (success, user_message_or_none)."""
    from extensions import db
    from models import Agency
    from services.gemini_letter import (
        build_fallback_complaint_letter,
        format_gemini_error,
        generate_complaint_letter,
    )

    agency = Agency.query.get(complaint.agency_id) if complaint.agency_id else None
    try:
        letter_text = generate_complaint_letter(complaint, complainant, agency)
        complaint.generated_letter = letter_text
        complaint.letter_generated = True
        db.session.commit()
        return True, None
    except Exception as exc:
        print(f'Gemini letter generation failed: {exc}')
        print(format_gemini_error(exc))
        complaint.generated_letter = build_fallback_complaint_letter(
            complaint,
            complainant,
            agency,
        )
        complaint.letter_generated = True
        db.session.commit()
        return True, 'fallback'


@user_bp.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        form = _register_form_from_request()
        full_name = form['full_name']
        email = form['email']
        contact_number = form['contact_number']
        barangay = form['barangay']
        municipality = form['municipality']
        password = form['password'].strip()
        confirm_password = form['confirm_password'].strip()

        if not all([full_name, email, contact_number, barangay, municipality, password, confirm_password]):
            flash('All fields are required.', 'danger')
            return _render_register_form(form)

        if not contact_number.isdigit() or len(contact_number) != 11:
            flash('Contact number must contain numbers only and be exactly 11 digits.', 'danger')
            return _render_register_form(form)

        if password != confirm_password:
            flash('Passwords do not match.', 'danger')
            return _render_register_form(form)

        password_error = _password_error(password)
        if password_error:
            flash(password_error, 'danger')
            return _render_register_form(form)

from flask import Blueprint, render_template, redirect, url_for, flash, request
from werkzeug.security import generate_password_hash, check_password_hash
from flask_login import login_user, logout_user, login_required, current_user
import re

from routes.auth import member_required

user_bp = Blueprint('user', __name__, url_prefix='/user')


# ── Register ──
@user_bp.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        full_name = request.form.get('full_name', '').strip()
        email = request.form.get('email', '').strip()
        contact_number = request.form.get('contact_number', '').strip()
        barangay = request.form.get('barangay', '').strip()
        municipality = request.form.get('municipality', '').strip()
        password = request.form.get('password', '').strip()
        confirm_password = request.form.get('confirm_password', '').strip()

        if not all([full_name, email, contact_number, barangay, municipality, password, confirm_password]):
            flash('All fields are required.', 'danger')
            return render_template('user/register.html')

        if not contact_number.isdigit() or len(contact_number) != 11:
            flash('Contact number must be exactly 11 digits.', 'danger')
            return render_template('user/register.html')

        if password != confirm_password:
            flash('Passwords do not match.', 'danger')
            return render_template('user/register.html')

        if len(password) < 8:
            flash('Password must be at least 8 characters.', 'danger')
            return render_template('user/register.html')

        if not re.search(r'[A-Z]', password):
            flash('Password must contain at least one uppercase letter.', 'danger')
            return render_template('user/register.html')

        if not re.search(r'[a-z]', password):
            flash('Password must contain at least one lowercase letter.', 'danger')
            return render_template('user/register.html')

        if not re.search(r'[0-9]', password):
            flash('Password must contain at least one number.', 'danger')
            return render_template('user/register.html')

        if not re.search(r'[!@#$%^&*(),.?":{}|<>]', password):
            flash('Password must contain at least one special character.', 'danger')
            return render_template('user/register.html')

        from models import User
        from extensions import db
        existing_user = User.query.filter_by(email=email).first()
        if existing_user:
            if existing_user.status == 'pending':
                flash(
                    'This email is already registered but not yet verified. '
                    'Please check your email for the OTP or request a new code.',
                    'danger',
                )
            else:
                flash('Email is already registered.', 'danger')
            return _render_register_form(form)

            flash('Email is already registered.', 'danger')
            return render_template('user/register.html')

        new_user = User(
            full_name=full_name,
            email=email,
            contact_number=contact_number,
            barangay=barangay,
            municipality=municipality,
            status='pending',
            password_hash=generate_password_hash(password)
        )
        db.session.add(new_user)
        db.session.commit()

        verification = _create_email_verification(new_user)
        sent = _send_account_verification_email(new_user, verification.otp_code)

        if sent:
            flash(
                'Account created successfully! An OTP has been sent to your email. '
                'Please verify your account to continue.',
                'success',
            )
        else:
            flash(
                'Account created successfully, but we could not send the verification email. '
                'Please contact support or try registering again later.',
                'warning',
            )

        return redirect(url_for('user.verify_account', user_id=new_user.id))

    return _render_register_form()


@user_bp.route('/verify/<int:user_id>', methods=['GET', 'POST'])
def verify_account(user_id):
    from extensions import db
    from models import EmailVerification, User

    user = User.query.get_or_404(user_id)
    if user.status == 'active':
        flash('Your account is already verified. Please log in.', 'info')
        return redirect(url_for('user.user_login'))

    if request.method == 'POST':
        otp = request.form.get('otp', '').strip()
        if not otp:
            flash('Please enter the OTP sent to your email.', 'danger')
            # compute remaining cooldown for template
            latest = EmailVerification.query.filter_by(user_id=user.id).order_by(EmailVerification.created_at.desc()).first()
            remaining = 0
            if latest:
                remaining = max(0, int((latest.created_at + timedelta(seconds=RESEND_COOLDOWN_SECONDS) - datetime.utcnow()).total_seconds()))
            return render_template('user/verify_account.html', user=user, resend_cooldown=remaining)

        verification = EmailVerification.query.filter_by(
            user_id=user.id,
            otp_code=otp,
            is_used=False,
        ).order_by(EmailVerification.created_at.desc()).first()

        if not verification or verification.expires_at < datetime.utcnow():
            flash('The OTP is invalid or has expired. Please request a new code.', 'danger')
            latest = EmailVerification.query.filter_by(user_id=user.id).order_by(EmailVerification.created_at.desc()).first()
            remaining = 0
            if latest:
                remaining = max(0, int((latest.created_at + timedelta(seconds=RESEND_COOLDOWN_SECONDS) - datetime.utcnow()).total_seconds()))
            return render_template('user/verify_account.html', user=user, resend_cooldown=remaining)

        verification.is_used = True
        user.status = 'active'
        db.session.commit()

        flash('Your account has been verified! Please log in.', 'success')
        return redirect(url_for('user.user_login'))

    latest = EmailVerification.query.filter_by(user_id=user.id).order_by(EmailVerification.created_at.desc()).first()
    remaining = 0
    if latest:
        remaining = max(0, int((latest.created_at + timedelta(seconds=RESEND_COOLDOWN_SECONDS) - datetime.utcnow()).total_seconds()))
    return render_template('user/verify_account.html', user=user, resend_cooldown=remaining)


@user_bp.route('/verify/<int:user_id>/resend', methods=['POST'])
def resend_verification_otp(user_id):
    from models import User

    user = User.query.get_or_404(user_id)
    if user.status == 'active':
        flash('Your account is already verified. Please log in.', 'info')
        return redirect(url_for('user.user_login'))
    from models import EmailVerification
    latest = EmailVerification.query.filter_by(user_id=user.id).order_by(EmailVerification.created_at.desc()).first()
    if latest:
        elapsed = (datetime.utcnow() - latest.created_at).total_seconds()
        if elapsed < RESEND_COOLDOWN_SECONDS:
            wait = int(RESEND_COOLDOWN_SECONDS - elapsed)
            flash(f'Please wait {wait} seconds before requesting a new code.', 'warning')
            return redirect(url_for('user.verify_account', user_id=user.id))

    verification = _create_email_verification(user)
    sent = _send_account_verification_email(user, verification.otp_code)
    if sent:
        flash('A new OTP has been sent to your email address.', 'success')
    else:
        flash(
            'Unable to send a new OTP right now. Please try again later or contact support.',
            'warning',
        )
    return redirect(url_for('user.verify_account', user_id=user.id))



    flash('Account created successfully! Please log in.', 'success')
    return redirect(url_for('user.user_login'))

    return render_template('user/register.html')


# ── User Login ──
@user_bp.route('/login', methods=['GET', 'POST'])
def user_login():
    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '').strip()

        if not email or not password:
            flash('All fields are required.', 'danger')
            return render_template('user/login.html')

        from models import User
        user = User.query.filter_by(email=email).first()

        if not user or not verify_password(user.password_hash, password):
            flash('Invalid credentials. Please try again.', 'danger')
            return render_template('user/login.html')

        if user.status == 'pending':
            flash(
                'Your account is not yet verified. Please check your email for the OTP ',
                'danger',
            )
            return render_template('user/login.html')

        if user.status == 'suspended':
            flash('Your account has been suspended. Please contact the administrator.', 'danger')
            return render_template('user/login.html')


        # Re-hash password with bcrypt if stored hash is not bcrypt-based
        try:
            if user and not is_bcrypt_hash(user.password_hash):
                user.password_hash = hash_password(password)
                from extensions import db
                db.session.commit()
        except Exception:
            pass


        #  NEW BCRYPT CHECK WHICH WORKS WITH YOUR UTILS FILE
        if not user or not verify_password(user.password_hash, password):
            flash('Invalid credentials. Please try again.', 'danger')
            return render_template('user/login.html')

        if user.status == 'suspended':
            flash('Your account has been suspended. Please contact admin.', 'danger')
            return render_template('user/login.html')

        login_user(user)
        return redirect(url_for('user.submit_complaint'))

    return render_template('user/login.html')


@user_bp.route('/logout')
@member_required

# ── User Logout ──
def user_logout():
    logout_user()
    return redirect(url_for('user.user_login'))



# ── Forgot Password ──
from itsdangerous import URLSafeTimedSerializer
from flask import current_app


def generate_reset_token(email):
    s = URLSafeTimedSerializer(current_app.config['SECRET_KEY'])
    return s.dumps(email, salt='password-reset-salt')


def verify_reset_token(token, expiration=3600):
    s = URLSafeTimedSerializer(current_app.config['SECRET_KEY'])
    try:
        email = s.loads(token, salt='password-reset-salt', max_age=expiration)
    except Exception:
        return None
    return email


@user_bp.route('/forgot-password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        email = request.form.get('email', '').strip()

        if not email:
            flash('Please enter your email address.', 'danger')
            return render_template('user/forgot_password.html')

        from models import User
        user = User.query.filter_by(email=email).first()

        if user:
            token = generate_reset_token(email)
            reset_url = url_for('user.reset_password', token=token, _external=True)
            body = f"""
            <h3>INGAT — Password Reset</h3>
            <p>Hello {user.full_name},</p>
            <p>Click the link below to reset your password:</p>
            <a href="{reset_url}" style="background:#2D6A4F;color:white;padding:10px 20px;
            border-radius:8px;text-decoration:none;">Reset Password</a>
            <p>This link expires in <strong>1 hour</strong>.</p>
            <p>If you did not request this, ignore this email.</p>
            """
            from utils import send_email
            send_email(email, 'INGAT — Password Reset Request', body)

        flash('If that email is registered, a reset link has been sent.', 'info')
        return redirect(url_for('user.forgot_password'))

    return render_template('user/forgot_password.html')



def reset_password(token):
    email = verify_reset_token(token)
    if not email:
        flash('The reset link is invalid or has expired.', 'danger')
        return redirect(url_for('user.forgot_password'))

    if request.method == 'POST':
        password = request.form.get('password', '').strip()
        confirm_password = request.form.get('confirm_password', '').strip()

        if not password or not confirm_password:
            flash('All fields are required.', 'danger')
            return render_template('user/reset_password.html', token=token)

        if password != confirm_password:
            flash('Passwords do not match.', 'danger')
            return render_template('user/reset_password.html', token=token)

        if len(password) < 8:
            flash('Password must be at least 8 characters.', 'danger')
            return render_template('user/reset_password.html', token=token)

        from models import User
        from extensions import db
        user = User.query.filter_by(email=email).first()
        if user:
            user.password_hash = generate_password_hash(password)
            db.session.commit()
            flash('Password reset successfully! Please log in.', 'success')
            return redirect(url_for('user.user_login'))

    return render_template('user/reset_password.html', token=token)


# ── Submit Complaint ──
@user_bp.route('/submit', methods=['GET', 'POST'])
@member_required
def submit_complaint():
    if request.method == 'POST':
        form = _complaint_form_from_request()
        violation_type = form['violation_type']
        street_address = form['street_address']
        barangay = form['barangay']
        municipality = form['municipality']
        date_incident = form['date_incident']
        description = form['description'].strip()

        violation_type = request.form.get('violation_type', '').strip()
        street_address = request.form.get('street_address', '').strip()
        barangay = request.form.get('barangay', '').strip()
        municipality = request.form.get('municipality', '').strip()
        date_incident = request.form.get('date_incident', '').strip()
        description = request.form.get('description', '').strip()
        photo = request.files.get('photo')

        if not violation_type:
            flash('Please select a violation type.', 'danger')
            return _render_submit_form(form)

        if not street_address:
            flash('Please enter a street address.', 'danger')
            return _render_submit_form(form)

        if not barangay:
            flash('Please enter a barangay.', 'danger')
            return _render_submit_form(form)

        if not municipality:
            flash('Please enter a municipality.', 'danger')
            return _render_submit_form(form)

        if not date_incident:
            flash('Please select the date of incident.', 'danger')
            return _render_submit_form(form)

        if not description:
            flash('Please enter a description of the violation.', 'danger')
            return _render_submit_form(form)

        if len(description) < 20:
            flash(
                f'Description must be at least 20 characters (you entered {len(description)}).',
                'danger',
            )
            return _render_submit_form(form)

        if violation_type not in ALLOWED_VIOLATION_TYPES:
            flash('Invalid violation type selected.', 'danger')
            return _render_submit_form(form)

        from datetime import date
        try:
            incident_date = date.fromisoformat(date_incident)
            if incident_date > date.today():
                flash('Date of incident cannot be a future date.', 'danger')
                return _render_submit_form(form)
        except ValueError:
            flash('Invalid date format.', 'danger')
            return _render_submit_form(form)

            return render_template('user/submit_complaint.html')
        except ValueError:
            flash('Invalid date format.', 'danger')
            return render_template('user/submit_complaint.html')

        photo_path = None
        if photo and photo.filename != '':
            import os
            from werkzeug.utils import secure_filename

            from datetime import datetime as dt

            allowed_extensions = {'jpg', 'jpeg', 'png'}
            ext = photo.filename.rsplit('.', 1)[-1].lower()
            if ext not in allowed_extensions:
                flash('Photo must be JPG or PNG only.', 'danger')
                return _render_submit_form(form)

                return render_template('user/submit_complaint.html')

            photo.seek(0, 2)
            file_size = photo.tell()
            photo.seek(0)
            if file_size > 5 * 1024 * 1024:
                flash('Photo must not exceed 5MB.', 'danger')
                return _render_submit_form(form)

            from datetime import datetime

            filename = secure_filename(photo.filename)
            unique_name = f'{current_user.id}_{datetime.utcnow().strftime("%Y%m%d%H%M%S")}_{filename}'

            return render_template('user/submit_complaint.html')

            filename = secure_filename(photo.filename)
            unique_name = f'{current_user.id}_{dt.utcnow().strftime("%Y%m%d%H%M%S")}_{filename}'

            upload_folder = os.path.join('static', 'uploads')
            os.makedirs(upload_folder, exist_ok=True)
            photo_path = os.path.join(upload_folder, unique_name)
            photo.save(photo_path)

        agency_map = {
            'Illegal Dumping': 'LGU',
            'Air Pollution': 'DENR',
            'Water Pollution': 'LLDA',
            'Illegal Logging': 'DENR',
            'Others': 'LGU'
        }
        agency_name = agency_map.get(violation_type, 'LGU')

        from models import Agency, Complaint
        from extensions import db

        agency = Agency.query.filter_by(agency_name=agency_name).first()
        agency_id = agency.id if agency else None

        new_complaint = Complaint(
            user_id=current_user.id,
            agency_id=agency_id,
            violation_type=violation_type,
            street_address=street_address,
            barangay=barangay,
            municipality=municipality,
            date_incident=incident_date,
            description=description,
            photo_path=photo_path,
            status='Submitted'
        )
        db.session.add(new_complaint)
        db.session.commit()

        letter_ok, letter_error = _generate_letter_for_complaint(new_complaint, current_user)

        if letter_error == 'fallback':
            flash(
                'Complaint submitted successfully. A formal letter was generated using the offline template.',
                'info',
            )
        elif letter_ok:
            flash('Complaint submitted successfully!', 'success')
        else:
            flash(f'Complaint saved. {letter_error}', 'warning')

        return redirect(url_for('user.complaint_submitted', complaint_id=new_complaint.id))

    return _render_submit_form()


@user_bp.route('/submitted/<int:complaint_id>')
@member_required
def complaint_submitted(complaint_id):
    from models import Complaint

    complaint = Complaint.query.get_or_404(complaint_id)
    if complaint.user_id != current_user.id:
        abort(403)
    return render_template('user/complaint_submitted.html', complaint=complaint)


def _letter_redirect(complaint_id, from_report=False):
    if from_report:
        return redirect(url_for('user.report_detail', complaint_id=complaint_id))
    return redirect(url_for('user.complaint_submitted', complaint_id=complaint_id))


@user_bp.route('/submitted/<int:complaint_id>/regenerate-letter', methods=['POST'])
@member_required
def regenerate_letter(complaint_id):
    from models import Complaint

    complaint = Complaint.query.get_or_404(complaint_id)
    if complaint.user_id != current_user.id:
        abort(403)

    from_report = request.form.get('from') == 'report'

    if complaint.letter_generated and complaint.generated_letter:
        flash('Letter is already available for this complaint.', 'info')
        return _letter_redirect(complaint_id, from_report)

    letter_ok, letter_error = _generate_letter_for_complaint(complaint, current_user)
    if letter_ok:
        flash('Formal complaint letter generated successfully!', 'success')
    else:
        flash(letter_error, 'warning')

    return _letter_redirect(complaint_id, from_report)


def _get_owned_complaint_with_letter(complaint_id):
    from models import Complaint

    complaint = Complaint.query.get_or_404(complaint_id)
    if complaint.user_id != current_user.id:
        abort(403)
    if not complaint.letter_generated or not complaint.generated_letter:
        flash('No generated letter is available for this complaint.', 'warning')
        return None
    return complaint


@user_bp.route('/submitted/<int:complaint_id>/download/pdf')
@member_required
def download_letter_pdf(complaint_id):
    complaint = _get_owned_complaint_with_letter(complaint_id)
    if not complaint:
        return redirect(url_for('user.complaint_submitted', complaint_id=complaint_id))

    from services.letter_export import build_letter_pdf

    pdf_buffer = build_letter_pdf(complaint.generated_letter, complaint.id)
    filename = f'INGAT_Complaint_{complaint.id:04d}.pdf'
    return send_file(pdf_buffer, mimetype='application/pdf', as_attachment=True, download_name=filename)


@user_bp.route('/submitted/<int:complaint_id>/download/docx')
@member_required
def download_letter_docx(complaint_id):
    complaint = _get_owned_complaint_with_letter(complaint_id)
    if not complaint:
        return redirect(url_for('user.complaint_submitted', complaint_id=complaint_id))

    from services.letter_export import build_letter_docx

    docx_buffer = build_letter_docx(complaint.generated_letter, complaint.id)
    filename = f'INGAT_Complaint_{complaint.id:04d}.docx'
    return send_file(
        docx_buffer,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename,
    )


def forgot_password():
    if request.method == 'POST':
        email = request.form.get('email', '').strip()

        if not email:
            flash('Please enter your email address.', 'danger')
            return render_template('user/forgot_password.html')

        from models import User
        user = User.query.filter_by(email=email).first()

        if user:
            token = generate_reset_token(email)
            reset_url = url_for('user.reset_password', token=token, _external=True)
            body = f"""
            <h3>INGAT — Password Reset</h3>
            <p>Hello {user.full_name},</p>
            <p>Click the link below to reset your password:</p>
            <a href="{reset_url}" style="background:#2D6A4F;color:white;padding:10px 20px;
            border-radius:8px;text-decoration:none;">Reset Password</a>
            <p>This link expires in <strong>1 hour</strong>.</p>
            <p>If you did not request this, ignore this email.</p>
            """
            send_email(email, 'INGAT — Password Reset Request', body)

        flash('If that email is registered, a reset link has been sent.', 'info')
        return redirect(url_for('user.forgot_password'))

    return render_template('user/forgot_password.html')


@user_bp.route('/reset-password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    email = verify_reset_token(token)
    if not email:
        flash('The reset link is invalid or has expired.', 'danger')
        return redirect(url_for('user.forgot_password'))

    if request.method == 'POST':
        password = request.form.get('password', '').strip()
        confirm_password = request.form.get('confirm_password', '').strip()

        if not password or not confirm_password:
            flash('All fields are required.', 'danger')
            return render_template('user/reset_password.html', token=token)

        if password != confirm_password:
            flash('Passwords do not match.', 'danger')
            return render_template('user/reset_password.html', token=token)

        password_error = _password_error(password)
        if password_error:
            flash(password_error, 'danger')
            return render_template('user/reset_password.html', token=token)

        from models import User
        from extensions import db
        user = User.query.filter_by(email=email).first()
        if user:
            user.password_hash = hash_password(password)
            db.session.commit()
            flash('Password reset successfully! Please log in.', 'success')
            return redirect(url_for('user.user_login'))

    return render_template('user/reset_password.html', token=token)


def _complaint_photo_url(complaint):
    if not complaint.photo_path:
        return None
    path = complaint.photo_path.replace('\\', '/')
    if path.startswith('static/'):
        path = path[len('static/'):]
    return url_for('static', filename=path)


@user_bp.route('/my-reports')
@member_required
def my_reports():
    from models import Complaint

    status_filter = request.args.get('status', '').strip()
    violation_filter = request.args.get('violation_type', '').strip()
    search = request.args.get('q', '').strip()

    query = Complaint.query.filter_by(user_id=current_user.id)
    if status_filter:
        query = query.filter_by(status=status_filter)
    if violation_filter:
        query = query.filter_by(violation_type=violation_filter)
    if search:
        normalized = search.upper().replace('#ING-', '').replace('ING-', '').strip()
        try:
            query = query.filter_by(id=int(normalized))
        except ValueError:
            query = query.filter(Complaint.id == -1)

    complaints = query.order_by(Complaint.created_at.desc()).all()
    all_statuses = ['Submitted', 'Under Review', 'Forwarded to Agency', 'Resolved']

    return render_template(
        'user/my_reports.html',
        complaints=complaints,
        status_filter=status_filter,
        violation_filter=violation_filter,
        search=search,
        all_statuses=all_statuses,
        violation_types=ALLOWED_VIOLATION_TYPES,
    )


@user_bp.route('/my-reports/<int:complaint_id>')
@member_required
def report_detail(complaint_id):
    from models import Complaint

    complaint = Complaint.query.get_or_404(complaint_id)
    if complaint.user_id != current_user.id:
        abort(403)

    status_history = sorted(
        complaint.status_history,
        key=lambda entry: entry.updated_at,
        reverse=True,
    )
    return render_template(
        'user/report_detail.html',
        complaint=complaint,
        status_history=status_history,
        photo_url=_complaint_photo_url(complaint),
    )

        # Generate AI letter
    from utils import generate_complaint_letter
    letter = generate_complaint_letter(
            violation_type=violation_type,
            description=description,
            barangay=barangay,
            municipality=municipality,
            date_incident=date_incident,
            complainant_name=current_user.full_name,
            contact_number=current_user.contact_number,
            agency_name=agency_name
        )

    if letter:
            new_complaint.generated_letter = letter
            new_complaint.letter_generated = True
            db.session.commit()
            flash('Complaint submitted successfully!', 'success')
    else:
            flash('Complaint saved. Letter generation failed.', 'warning')

    return redirect(url_for('user.complaint_submitted', complaint_id=new_complaint.id))

    return render_template('user/submit_complaint.html')


# ── Complaint Submitted ──
@login_required
def complaint_submitted(complaint_id):
    from models import Complaint
    complaint = Complaint.query.get_or_404(complaint_id)
    if complaint.user_id != current_user.id:
        flash('Unauthorized access.', 'danger')
        return redirect(url_for('user.my_reports'))
    return render_template('user/complaint_submitted.html', complaint=complaint)


# ── My Reports ──
@login_required
def my_reports():
    return render_template('user/my_reports.html')


# ── Download PDF
@login_required
def download_letter_pdf(complaint_id):
    from models import Complaint
    from fpdf import FPDF
    from flask import make_response

    complaint = Complaint.query.get_or_404(complaint_id)

    if complaint.user_id != current_user.id:
        flash('Unauthorized access.', 'danger')
        return redirect(url_for('user.my_reports'))

    pdf = FPDF()
    pdf.add_page()
    pdf.set_margins(20, 20, 20)
    pdf.set_font('Helvetica', 'B', 14)
    pdf.cell(0, 10, 'Formal Complaint Letter', ln=True, align='C')

    pdf.set_font('Helvetica', size=11)
    pdf.cell(0, 8, f'Reference: #ING-{complaint.id:04d}', ln=True, align='C')
    pdf.ln(8)

    letter_text = complaint.generated_letter or 'No letter generated.'

    # Robust rendering to avoid FPDF multi_cell crashes
    for raw_line in letter_text.split('\n'):
        safe_line = (raw_line or '').encode('latin-1', 'replace').decode('latin-1')
        try:
            pdf.multi_cell(0, 7, safe_line)
        except Exception:
            for ch in safe_line:
                try:
                    pdf.cell(0, 7, ch, ln=1)
                except Exception:
                    pdf.cell(0, 7, '?', ln=1)

    pdf_bytes = pdf.output(dest='S')
    response = make_response(pdf_bytes.encode('latin-1') if isinstance(pdf_bytes, str) else pdf_bytes)
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = f'attachment; filename=INGAT-Complaint-{complaint.id:04d}.pdf'
    return response


# ── Download DOCX ──
@login_required
def download_letter_docx(complaint_id):
    from models import Complaint
    from docx import Document
    import io
    from flask import make_response

    complaint = Complaint.query.get_or_404(complaint_id)

    if complaint.user_id != current_user.id:
        flash('Unauthorized access.', 'danger')
        return redirect(url_for('user.my_reports'))

    doc = Document()
    title = doc.add_heading('Formal Complaint Letter', 0)
    title.alignment = 1
    ref = doc.add_paragraph(f'Reference: #ING-{complaint.id:04d}')
    ref.alignment = 1
    doc.add_paragraph('')

    letter_text = complaint.generated_letter or 'No letter generated.'
    for line in letter_text.split('\n'):
        doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = make_response(buffer.read())
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    response.headers['Content-Disposition'] = f'attachment; filename=INGAT-Complaint-{complaint.id:04d}.docx'
    return response

